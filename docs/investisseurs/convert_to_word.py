#!/usr/bin/env python3
"""
Script amélioré pour convertir les documents Markdown en fichiers Word (.docx)
avec formatage professionnel et correction des problèmes d'exportation
"""

import os
import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installation de python-docx...")
    os.system(f"{sys.executable} -m pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# Mapping des emojis vers texte
EMOJI_MAP = {
    '📊': '', '🎯': '', '💡': '', '🚀': '', '📈': '', '💰': '', '👥': '',
    '⚙️': '', '📢': '', '🗺️': '', '📱': '', '🎨': '', '🌍': '', '🚨': '',
    '✅': '', '❌': '', '⚠️': '', '📋': '', '📍': '', '🏢': '', '💼': '',
    '📅': '', '🔄': '', '📊': '', '🎯': '', '💸': '', '📚': '', '🔍': ''
}

def clean_text(text):
    """Nettoie le texte des emojis et caractères problématiques"""
    # Supprimer les emojis
    for emoji, replacement in EMOJI_MAP.items():
        text = text.replace(emoji, replacement)
    
    # Supprimer les étoiles isolées qui restent (mais pas celles dans **text**)
    # D'abord protéger les patterns **text**
    protected_patterns = []
    pattern = r'\*\*([^*]+)\*\*'
    for i, match in enumerate(re.finditer(pattern, text)):
        placeholder = f'__PROTECTED_{i}__'
        protected_patterns.append((placeholder, match.group(0)))
        text = text.replace(match.group(0), placeholder, 1)
    
    # Supprimer les étoiles isolées
    text = re.sub(r'\s+\*\s+', ' ', text)
    text = re.sub(r'^\*\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\s+\*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\s+', '', text)
    text = re.sub(r'\s+\*', '', text)
    
    # Restaurer les patterns protégés
    for placeholder, original in protected_patterns:
        text = text.replace(placeholder, original)
    
    return text.strip()

def process_bold_text(text):
    """Traite le texte en gras (**text**) et retourne les parties"""
    parts = []
    pattern = r'\*\*([^*]+)\*\*'
    
    last_end = 0
    for match in re.finditer(pattern, text):
        # Ajouter le texte avant le match
        if match.start() > last_end:
            parts.append(('normal', text[last_end:match.start()]))
        
        # Ajouter le texte en gras
        parts.append(('bold', match.group(1)))
        last_end = match.end()
    
    # Ajouter le texte restant
    if last_end < len(text):
        parts.append(('normal', text[last_end:]))
    
    if not parts:
        parts.append(('normal', text))
    
    return parts

def process_inline_code(text):
    """Traite le code inline (`code`)"""
    parts = []
    pattern = r'`([^`]+)`'
    
    last_end = 0
    for match in re.finditer(pattern, text):
        if match.start() > last_end:
            parts.append(('normal', text[last_end:match.start()]))
        parts.append(('code', match.group(1)))
        last_end = match.end()
    
    if last_end < len(text):
        parts.append(('normal', text[last_end:]))
    
    if not parts:
        parts.append(('normal', text))
    
    return parts

def process_references(text):
    """Supprime ou remplace les références [^1]"""
    # Supprimer les références de type [^1]
    text = re.sub(r'\[\^\d+\]', '', text)
    # Supprimer les références de type [^1]: en début de ligne
    text = re.sub(r'^\[\^\d+\]:\s*', '', text, flags=re.MULTILINE)
    return text

def add_formatted_paragraph(doc, text, style='Normal'):
    """Ajoute un paragraphe avec formatage avancé"""
    if not text or not text.strip():
        return doc.add_paragraph()
    
    # Nettoyer le texte
    text = clean_text(text)
    text = process_references(text)
    
    # Traiter d'abord le code inline, puis le gras
    p = doc.add_paragraph(style=style)
    
    # Traiter code inline d'abord
    code_parts = process_inline_code(text)
    
    for code_type, code_text in code_parts:
        if code_type == 'code':
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 0, 128)  # Violet pour code
        else:
            # Traiter le gras dans le texte normal
            bold_parts = process_bold_text(code_text)
            for bold_type, bold_text in bold_parts:
                if bold_text:
                    run = p.add_run(bold_text)
                    if bold_type == 'bold':
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 51, 102)  # Bleu foncé pour gras
    
    return p

def create_professional_table(doc, rows_data):
    """Crée un tableau professionnel"""
    if not rows_data or not rows_data[0]:
        return None
    
    num_cols = len(rows_data[0])
    num_rows = len(rows_data)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # Mettre en forme les cellules
    for row_idx, row_data in enumerate(rows_data):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.rows[row_idx].cells[col_idx]
                cell_text = clean_text(cell_text)
                cell_text = process_references(cell_text)
                
                # Nettoyer les pipes restants
                cell_text = cell_text.replace('|', '').strip()
                
                p = cell.paragraphs[0]
                p.clear()
                
                # Traiter le formatage
                parts = process_bold_text(cell_text)
                for part_type, part_text in parts:
                    if part_text:
                        run = p.add_run(part_text)
                        if part_type == 'bold':
                            run.bold = True
                            run.font.color.rgb = RGBColor(0, 51, 102)
                
                # Style en-tête (première ligne)
                if row_idx == 0:
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(11)
                    # Fond gris pour en-tête
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'E7E6E6')
                    cell._element.get_or_add_tcPr().append(shading_elm)
                else:
                    p.runs[0].font.size = Pt(10)
    
    return table

def markdown_to_docx(md_file, docx_file):
    """Convertit un fichier Markdown en fichier Word avec formatage professionnel"""
    
    # Lire le fichier Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Créer un nouveau document Word
    doc = Document()
    
    # Configurer les styles par défaut
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    
    # Configurer les styles de titres
    for i in range(1, 5):
        heading_style = doc.styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = 'Calibri'
        if i == 1:
            heading_font.size = Pt(18)
            heading_font.bold = True
            heading_font.color.rgb = RGBColor(0, 51, 102)
        elif i == 2:
            heading_font.size = Pt(16)
            heading_font.bold = True
            heading_font.color.rgb = RGBColor(0, 51, 102)
        elif i == 3:
            heading_font.size = Pt(14)
            heading_font.bold = True
            heading_font.color.rgb = RGBColor(0, 51, 102)
        else:
            heading_font.size = Pt(12)
            heading_font.bold = True
    
    # Parser le Markdown ligne par ligne
    lines = md_content.split('\n')
    i = 0
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Ignorer les lignes vides multiples
        if not line:
            if i < len(lines) - 1 and lines[i+1].strip():
                doc.add_paragraph()
            i += 1
            continue
        
        # Titre niveau 1 (#)
        if line.startswith('# '):
            text = clean_text(line[2:].strip())
            text = process_references(text)
            doc.add_heading(text, level=1)
            i += 1
        
        # Titre niveau 2 (##)
        elif line.startswith('## '):
            text = clean_text(line[3:].strip())
            text = process_references(text)
            doc.add_heading(text, level=2)
            i += 1
        
        # Titre niveau 3 (###)
        elif line.startswith('### '):
            text = clean_text(line[4:].strip())
            text = process_references(text)
            doc.add_heading(text, level=3)
            i += 1
        
        # Titre niveau 4 (####)
        elif line.startswith('#### '):
            text = clean_text(line[5:].strip())
            text = process_references(text)
            doc.add_heading(text, level=4)
            i += 1
        
        # Ligne horizontale (---)
        elif line.startswith('---'):
            p = doc.add_paragraph()
            p.add_run('_' * 80).font.color.rgb = RGBColor(200, 200, 200)
            i += 1
        
        # Tableau Markdown
        elif '|' in line and line.count('|') >= 2:
            if not in_table:
                in_table = True
                table_lines = []
            
            # Ignorer les lignes de séparation
            if not line.startswith('|---'):
                table_lines.append(line)
            i += 1
        
        # Fin de tableau
        elif in_table and '|' not in line:
            if table_lines:
                rows_data = []
                for tl in table_lines:
                    cells = [c.strip() for c in tl.split('|') if c.strip()]
                    if cells:
                        rows_data.append(cells)
                
                if rows_data:
                    create_professional_table(doc, rows_data)
            
            in_table = False
            table_lines = []
            # Ne pas incrémenter i, traiter la ligne actuelle
        
        # Liste à puces (- ou *)
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            text = clean_text(text)
            text = process_references(text)
            
            # Vérifier le niveau d'indentation
            indent_level = 0
            if line.startswith('  - ') or line.startswith('  * '):
                indent_level = 1
            elif line.startswith('    - ') or line.startswith('    * '):
                indent_level = 2
            
            p = doc.add_paragraph(style='List Bullet' if indent_level == 0 else 'List Bullet 2')
            
            # Traiter le formatage
            parts = process_bold_text(text)
            for part_type, part_text in parts:
                if part_text:
                    run = p.add_run(part_text)
                    if part_type == 'bold':
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 51, 102)
            
            i += 1
        
        # Paragraphe normal
        else:
            if not in_table:
                add_formatted_paragraph(doc, line)
            i += 1
    
    # Traiter le dernier tableau si on est encore dedans
    if in_table and table_lines:
        rows_data = []
        for tl in table_lines:
            cells = [c.strip() for c in tl.split('|') if c.strip()]
            if cells:
                rows_data.append(cells)
        
        if rows_data:
            create_professional_table(doc, rows_data)
    
    # Ajouter les en-têtes et pieds de page
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_run = header_para.add_run('Yukpomnang - Document Investisseurs')
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = RGBColor(128, 128, 128)
    
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('Confidentiel - Janvier 2026')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Sauvegarder le document
    try:
        doc.save(docx_file)
        print(f"[OK] Converti : {Path(md_file).name} -> {Path(docx_file).name}")
    except PermissionError:
        print(f"[ERREUR] Permission refusee : {Path(docx_file).name} est peut-etre ouvert dans Word")
        print(f"         Veuillez fermer le fichier et reessayer")
    except Exception as e:
        print(f"[ERREUR] Erreur lors de la conversion de {md_file}: {e}")

def main():
    """Fonction principale"""
    current_dir = Path(__file__).parent
    
    # Liste des fichiers à convertir
    md_files = [
        '01_EXECUTIVE_SUMMARY.md',
        '02_BUSINESS_PLAN.md',
        '03_PROJECTIONS_FINANCIERES.md',
        '04_PLAN_MARKETING.md',
        '05_ANALYSE_MARCHE.md',
        '06_PLAN_OPERATIONNEL.md',
        '07_PITCH_DECK.md',
        '08_PLAN_FINANCIER_3_ANS.md',
    ]
    
    print("=" * 60)
    print("Conversion des documents Markdown en Word")
    print("Formatage professionnel ameliore")
    print("=" * 60)
    print()
    
    success_count = 0
    error_count = 0
    
    for md_file in md_files:
        md_path = current_dir / md_file
        if md_path.exists():
            docx_file = md_path.stem + '.docx'
            docx_path = current_dir / docx_file
            try:
                markdown_to_docx(md_path, docx_path)
                success_count += 1
            except Exception as e:
                print(f"[ERREUR] Erreur lors de la conversion de {md_file}: {e}")
                error_count += 1
        else:
            print(f"[ATTENTION] Fichier non trouve : {md_file}")
            error_count += 1
    
    print()
    print("=" * 60)
    print(f"Conversion terminée : {success_count} succès, {error_count} erreurs")
    print("=" * 60)

if __name__ == '__main__':
    main()
